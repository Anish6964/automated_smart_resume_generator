from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
from fuzzywuzzy import process  # For intelligent matching


# Helper function to set font style
def set_font(run, font_name="Calibri", size=Pt(11), bold=False, color=None):
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(color[0], color[1], color[2])  # RGB color value


# Helper function to set cell style in a table
def set_cell_style(
    cell,
    text,
    font_name="Calibri",
    size=Pt(11),
    bold=False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    bg_color=None,
    text_color=None,
):
    # Set text inside the cell
    cell_paragraph = cell.paragraphs[0]
    run = cell_paragraph.add_run(text)
    set_font(run, font_name, size, bold, text_color)
    cell_paragraph.alignment = alignment
    set_cell_border(cell)

    # Set background color of the cell
    if bg_color:
        set_cell_background_color(cell, bg_color)


# Helper function to set table cell borders
def set_cell_border(cell):
    # Get the table cell properties
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    # Create the borders element
    borders = OxmlElement("w:tcBorders")

    # Define each border side
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")  # Set single line border
        border.set(qn("w:sz"), "12")  # Border thickness (1/8 pt)
        border.set(qn("w:space"), "0")  # Space between text and border
        border.set(qn("w:color"), "000000")  # Border color in hex (black)
        borders.append(border)

    # Apply the borders to the cell
    tcPr.append(borders)


# Helper function to set cell background color using the w:shd element
def set_cell_background_color(cell, bg_color_hex):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    # Create the shading element (w:shd)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")  # Set to 'clear' for solid fill
    shd.set(qn("w:color"), "auto")  # Text color should be 'auto'
    shd.set(qn("w:fill"), bg_color_hex)  # Background color in hex (e.g., '4472C4')

    # Apply shading to the cell
    tcPr.append(shd)


# Helper function to find the best matching experience for a given responsibility
def find_best_match(responsibility, experiences):
    """
    This function takes a job responsibility and a list of experiences,
    and returns the experience that best matches the responsibility.
    """
    best_match, best_score = process.extractOne(responsibility, experiences)
    return best_match


# Function to create a more intelligent compatibility matrix
def create_intelligent_compatibility_matrix(user_info, job, output_dir):
    # Extract relevant experiences and skills to match against job responsibilities
    all_experiences = (
        user_info["key_achievements"]
        + [exp["results"] for exp in user_info["experience"]]
        + user_info["additional_skills"]
    )

    # Create a dictionary to store the best matching results
    compatibility_data = {}

    for responsibility in job["responsibilities"]:
        # Find the best match for each responsibility from the list of experiences and skills
        best_match = find_best_match(responsibility, all_experiences)
        compatibility_data[responsibility] = best_match

    # Now call the existing compatibility matrix creation function with the new data
    create_compatibility_matrix(user_info, job, compatibility_data, output_dir)


# Function to create a compatibility matrix document
def create_compatibility_matrix(user_info, job, compatibility_data, output_dir):
    # Create a new Document
    doc = Document()

    # Set the page margins to match the original template
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(72)  # Top margin
        section.bottom_margin = Pt(72)  # Bottom margin
        section.left_margin = Pt(72)  # Left margin
        section.right_margin = Pt(72)  # Right margin

    # Add title to the document
    title = doc.add_paragraph()
    title_run = title.add_run(f"Compatibility Matrix for {job['title']}")
    set_font(title_run, font_name="Calibri", size=Pt(14), bold=True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Create a table with two columns for Job Responsibilities and Demonstrated Results
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False

    # Set the table column widths
    table.columns[0].width = Pt(260)  # Job Responsibilities column width
    table.columns[1].width = Pt(260)  # Demonstrated Results column width

    # Create the table header
    header_cells = table.rows[0].cells
    set_cell_style(
        header_cells[0],
        "JOB RESPONSIBILITIES",
        font_name="Calibri",
        size=Pt(12),
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bg_color="4472C4",
        text_color=(255, 255, 255),
    )
    set_cell_style(
        header_cells[1],
        "DEMONSTRATED RESULTS",
        font_name="Calibri",
        size=Pt(12),
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bg_color="4472C4",
        text_color=(255, 255, 255),
    )

    # Populate the table with data from compatibility_data dictionary
    for responsibility, result in compatibility_data.items():
        # Create a single row with two cells: one for responsibility and one for result
        row_cells = table.add_row().cells

        # Set responsibility in the first cell as a single block of text (no breaks or splitting)
        set_cell_style(
            row_cells[0],
            responsibility,
            font_name="Calibri",
            size=Pt(11),
            bold=False,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            bg_color="1F4E78",
            text_color=(255, 255, 255),
        )

        # Set the entire result content in the second cell, formatted with bullet points for each line
        result_paragraph = row_cells[1].paragraphs[0]
        for line in result.split("\n"):
            # Add a bullet point for each line of result text
            run = result_paragraph.add_run(f"• {line.strip()}\n")
            set_font(
                run, font_name="Calibri", size=Pt(11), bold=False, color=(255, 255, 255)
            )
        result_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Apply the cell style and background color for the Demonstrated Results cell
        set_cell_background_color(
            row_cells[1], "2E2E2E"
        )  # Dark gray background for results cell
        set_cell_border(row_cells[1])  # Apply borders to result cells

    # Save the compatibility matrix in the specified output directory
    file_name = f"{user_info['name']}_Compatibility_Matrix_for_{job['title']}.docx"
    output_path = os.path.join(output_dir, file_name)
    doc.save(output_path)
    print(f"Compatibility matrix saved as {output_path}")


# Function to create a tailored resume
def create_resume(user_info, job, output_dir):
    doc = Document()

    # Header - Name and Contact Info
    header = doc.add_paragraph()
    name_run = header.add_run(user_info["name"] + "\n")
    set_font(
        name_run, font_name="Arial", size=Pt(16), bold=True
    )  # Larger font size for name

    contact_info = f"{user_info['location']} | {user_info['phone']} | {user_info['email']} | {user_info['linkedin']}"
    contact_run = header.add_run(contact_info)
    set_font(
        contact_run, font_name="Arial", size=Pt(11), bold=False
    )  # Regular font size for other details

    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Professional Summary Section
    if "summary" in user_info:
        add_paragraph(doc, "PROFESSIONAL SUMMARY", bold=True)
        add_paragraph(doc, user_info["summary"])

    # Key Achievements Section
    if "key_achievements" in user_info and user_info["key_achievements"]:
        add_paragraph(doc, "KEY ACHIEVEMENTS", bold=True)
        for achievement in user_info["key_achievements"]:
            add_paragraph(doc, f"• {achievement}")

    # Core Competencies Section
    if "core_competencies" in user_info and user_info["core_competencies"]:
        add_paragraph(doc, "CORE COMPETENCIES", bold=True)
        for competency in user_info["core_competencies"]:
            add_paragraph(doc, f"• {competency}")

    # Professional Experience Section
    if "experience" in user_info and user_info["experience"]:
        add_paragraph(doc, "PROFESSIONAL EXPERIENCE", bold=True)
        for experience in user_info["experience"]:
            add_paragraph(
                doc,
                f"{experience['title']}, {experience['company']}, {experience['location']} ({experience['dates']})",
                bold=True,
            )
            add_paragraph(doc, f"Skills: {experience['skills']}")
            add_paragraph(doc, f"Actions: {experience['actions']}")
            add_paragraph(doc, f"Results: {experience['results']}")

    # Education Section
    if "education" in user_info and user_info["education"]:
        add_paragraph(doc, "EDUCATION", bold=True)
        for education in user_info["education"]:
            add_paragraph(
                doc,
                f"{education['degree']}, {education['institution']}, {education['location']}",
            )

    # Professional Development Section
    if (
        "professional_development" in user_info
        and user_info["professional_development"]
    ):
        add_paragraph(doc, "PROFESSIONAL DEVELOPMENT", bold=True)
        for development in user_info["professional_development"]:
            add_paragraph(doc, f"• {development}")

    # Additional Skills and Achievements Section
    if "additional_skills" in user_info and user_info["additional_skills"]:
        add_paragraph(doc, "ADDITIONAL SKILLS AND ACHIEVEMENTS", bold=True)
        for skill in user_info["additional_skills"]:
            add_paragraph(doc, f"• {skill}")

    # Soft Skills Section
    if "soft_skills" in user_info and user_info["soft_skills"]:
        add_paragraph(doc, "SOFT SKILLS", bold=True)
        for soft_skill in user_info["soft_skills"]:
            add_paragraph(doc, f"• {soft_skill}")

    # Save the resume in the specified output directory
    file_name = os.path.join(
        output_dir, f"{user_info['name']}_Resume_for_{job['title']}.docx"
    )
    doc.save(file_name)
    print(f"Resume saved as {file_name}")


# Helper function to add paragraphs with consistent formatting
def add_paragraph(doc, text, bold=False):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_font(run, font_name="Arial", size=Pt(11), bold=bold)


# Function to create a tailored cover letter
def create_cover_letter(user_info, job, company_name, output_dir):
    doc = Document()

    # Introduction
    doc.add_paragraph(f"Dear Hiring Manager at {company_name}")

    # Body of the cover letter
    body_text = (
        f"I am excited to apply for the position of {job['title']}. Your company's mission to provide trusted medical solutions aligns "
        f"with my passion for improving patient outcomes through innovative healthcare solutions. As an RN with a deep understanding of clinical "
        f"practice, particularly in wound care, I have consistently driven sales growth and built meaningful relationships with healthcare professionals.\n\n"
        f"At Baxter International, I generated $4.6M in new revenue by understanding client needs and exceeding sales targets. I am confident that my experience "
        f"in advanced wound care products coupled with my consultative sales approach will allow me to drive success for your team.\n\n"
        f"I am eager to bring my skills and expertise to your team and contribute to the future growth of your organization. Thank you for considering my application. "
        f"I look forward to the opportunity to discuss how my background and experience align with your company's goals."
    )
    doc.add_paragraph(body_text)

    # Closing - "Sincerely," followed directly by contact information without a gap
    closing_paragraph = doc.add_paragraph()
    run = closing_paragraph.add_run("Sincerely,")
    set_font(run, font_name="Calibri", size=Pt(12), bold=False)
    closing_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Contact information in the same paragraph as Sincerely
    contact_info = f"\n{user_info['name']}\n{user_info['location']} | {user_info['phone']} | {user_info['email']} | {user_info['linkedin']}"
    run = closing_paragraph.add_run(contact_info)
    set_font(run, font_name="Calibri", size=Pt(12), bold=False)

    # Save the cover letter in the specified output directory
    file_name = f"{user_info['name']}_Cover_Letter_for_{job['title']}.docx"
    output_path = os.path.join(output_dir, file_name)
    doc.save(output_path)
    print(f"Cover letter saved as {file_name}")


# Function to process all job descriptions and generate documents for each
def process_all_jobs(user_info, job_descriptions, output_dir):
    for job in job_descriptions:
        # Create separate folders for each job application
        job_output_dir = os.path.join(output_dir, job["title"].replace(" ", "_"))
        os.makedirs(job_output_dir, exist_ok=True)

        # Create resume, cover letter, and compatibility matrix
        create_resume(user_info, job, job_output_dir)
        create_cover_letter(user_info, job, job["company"], job_output_dir)
        create_intelligent_compatibility_matrix(user_info, job, job_output_dir)


# Expanded user information with more details for stress testing
user_info = {
    "name": "Erin M. Briner, MHI, BSN, RN",
    "location": "Mpls, MN",
    "phone": "763.657.6906",
    "email": "embriner76@gmail.com",
    "linkedin": "https://www.linkedin.com/in/erin-briner/",
    "summary": (
        "Experienced Healthcare IT and Digital Health professional with over 15 years of experience in driving technology adoption, "
        "improving patient care outcomes, and exceeding sales targets. Proven track record in consultative sales, relationship management, "
        "and revenue growth across the healthcare industry."
    ),
    "key_achievements": [
        "Generated $4.6 million in net new revenue by forming strategic alliances and leveraging cross-selling opportunities in a seven-state region.",
        "Improved patient care outcomes by 30% through the implementation of advanced digital health platforms and real-time clinical dashboards.",
        "Exceeded sales targets by 105% in 2022 driving significant revenue growth through tailored sales strategies and strong relationship management.",
        "Increased customer satisfaction by 20% by enhancing operations and outcomes through expertise in IT hospital networks and clinical settings.",
        "Developed strong partnerships with essential clinical IT and C-suite stakeholders leading to a 25% increase in long-term contracts.",
        "Collaborated with cross-functional teams to streamline implementation processes, reducing onboarding time by 40%.",
        "Led the launch of a new SaaS product, achieving a 50% adoption rate within the first year of release.",
        "Developed a comprehensive training program for new hires, resulting in a 30% reduction in ramp-up time.",
        "Awarded 'Top Sales Performer' for two consecutive years for surpassing sales quotas and client engagement metrics.",
        "Established a new regional training program that led to a 50% increase in product knowledge and sales effectiveness.",
        "Created a data-driven sales strategy that increased territory revenue by 35% within the first year of implementation.",
    ],
    "core_competencies": [
        "Industry Knowledge: Expertise in the Healthcare Vertical trends, directions, major issues, regulatory considerations, and trendsetters.",
        "Effective Communications: Proficient in effective communication concepts, tools, and techniques.",
        "Negotiating: Skilled in successful negotiation concepts and techniques.",
        "Networking: Understanding the business value of creating mutually beneficial relationships.",
        "Cross-functional Collaboration: Knowledge of collaborative techniques and approaches.",
        "Prospecting: Proficient in prospecting principles, processes, and skills.",
        "Sales Closing and Agreements: Knowledgeable in sales closing and agreements processes, techniques, and skills.",
        "Strategic Sales Planning: Knowledge of sales principles, processes, techniques, and tools.",
        "Revenue Growth Strategies: Expertise in driving top-line growth through strategic initiatives and partnerships.",
        "Healthcare Informatics: Skilled in implementing digital health solutions to improve clinical outcomes.",
    ],
    "experience": [
        {
            "company": "BAXTER INTERNATIONAL",
            "title": "SaaS B2B Healthcare IT Sales Executive",
            "location": "Minneapolis, MN",
            "dates": "02/2021 - 01/2023",
            "skills": "Strategic Sales Planning, Healthcare Informatics, B2B SaaS Solutions, Relationship Management, Revenue Growth.",
            "actions": "Conducted market research, developed tailored sales strategies, fostered relationships with key stakeholders, and implemented advanced digital health platforms.",
            "results": "Generated $4.6M in net new revenue, exceeded 2022 sales targets by 105%, improved patient care outcomes.",
        },
        {
            "company": "AGILITI HEALTH",
            "title": "Regional Sales Manager",
            "location": "Minneapolis, MN",
            "dates": "01/2018 - 01/2021",
            "skills": "Team Leadership, Sales Management, Strategic Planning, Healthcare Solutions, Client Retention.",
            "actions": "Led a team of 8 account managers in achieving territory sales targets, managed key client accounts, and implemented customer retention strategies.",
            "results": "Achieved 120% of sales quota in 2019, increased client retention by 15%, and expanded market share by 10% through new business acquisition.",
        },
        {
            "company": "PHILIPS HEALTHCARE",
            "title": "Healthcare Consultant",
            "location": "Chicago, IL",
            "dates": "06/2014 - 12/2017",
            "skills": "Healthcare Consulting, Digital Health Solutions, Project Management, Clinical Operations.",
            "actions": "Provided consulting services to healthcare systems on digital health strategies, collaborated with clinical teams to identify operational improvements, and managed project implementation.",
            "results": "Led 5 successful implementations of clinical decision support systems, reducing adverse event rates by 20% and improving clinical workflow efficiency by 30%.",
        },
        # Additional experience entries for stress testing
        {
            "company": "HILLROM",
            "title": "Clinical Solutions Consultant",
            "location": "Chicago, IL",
            "dates": "02/2010 - 05/2014",
            "skills": "Clinical Workflow Optimization, Digital Health Adoption, Sales Strategy.",
            "actions": "Consulted with hospital staff on clinical workflow optimization and technology adoption strategies.",
            "results": "Increased adoption of digital health tools by 45%, improved staff efficiency by 25%, and contributed to a 15% reduction in patient stay duration.",
        },
    ],
    "education": [
        {
            "degree": "Master of Science (M.S.) in Health Informatics",
            "institution": "St. Catherine University",
            "location": "St. Paul, MN",
        },
        {
            "degree": "Bachelor of Science (B.S.) in Nursing",
            "institution": "St. Catherine University",
            "location": "St. Paul, MN",
        },
    ],
    "professional_development": [
        "Member of the Healthcare Information and Management Systems Society (HIMSS)",
        "Active participant in various hospital RN leadership groups",
        "Completed advanced training in SalesForce, Microsoft CRM, and other SaaS platforms",
        "Certification in Digital Health Strategy",
        "Attended the Annual Conference on Healthcare Innovation and Technology, 2021",
        "Certified in Lean Six Sigma methodologies for process optimization",
    ],
    "additional_skills": [
        "Technical Skills: Power BI, SugarCRM, SQL, Python, R, Tableau, Microsoft Word and Excel.",
        "Advanced Event Surveillance: Horizon Trends, Protocol Watch, Neonatal Event Review.",
        "Strategic Account Management: Led initiatives resulting in a 25% increase in revenue within six months.",
        "Expert in building relationships with clinical staff and hospital leadership to implement technology solutions.",
        "Experienced in leveraging data analytics to drive business decisions and optimize sales performance.",
    ],
    "soft_skills": [
        "Effective Communication",
        "Empathy",
        "Leadership",
        "Relationship Building",
        "Problem Solving",
        "Adaptability",
        "Critical Thinking",
        "Team Collaboration",
        "Customer-Centric Approach",
        "Time Management",
        "Conflict Resolution",
        "Cultural Sensitivity",
        "Negotiation",
        "Patience and Resilience",
        "Creativity",
        "Emotional Intelligence",
    ],
}

# Expanded job descriptions with more details for stress testing
job_descriptions = [
    {
        "title": "Territory Manager, Chronic Care - Minneapolis, MN",
        "company": "Convatec",
        "responsibilities": [
            "Developing and encouraging strong customer relationships and building brand loyalty.",
            "Driving current customer expansion through new product sales.",
            "Maintaining post-sales contact with large or strategic clients by facilitating a positive and productive long-term relationship.",
            "Building and maintaining relationships with key decision-makers that lead to future business opportunities.",
            "Use health economics data to develop territory business plans to meet/exceed assigned sales goals.",
            "Participate in training sessions to stay updated on product knowledge and healthcare trends.",
            "Prepare and deliver effective sales presentations to potential customers, addressing their pain points and demonstrating value.",
            "Track and analyze sales data to identify trends and optimize territory strategies.",
            "Work cross-functionally with marketing and operations teams to ensure seamless product delivery and support.",
        ],
    },
    {
        "title": "Senior Sales Executive - Digital Health Solutions",
        "company": "Philips Healthcare",
        "responsibilities": [
            "Identify new business opportunities in the healthcare technology sector and develop strategic sales plans.",
            "Collaborate with product management and marketing teams to develop go-to-market strategies for new product launches.",
            "Conduct in-depth market research to understand industry trends and competitor activities.",
            "Manage complex sales cycles from prospecting to closure, including contract negotiation and pricing discussions.",
            "Work closely with client stakeholders to understand their unique needs and propose tailored digital health solutions.",
            "Deliver high-impact presentations and demonstrations to showcase the value of products and services.",
            "Participate in industry conferences and events to build networks and identify potential partnership opportunities.",
            "Provide mentorship and coaching to junior sales representatives, fostering a culture of continuous learning and development.",
        ],
    },
    # Additional job descriptions for stress testing
    {
        "title": "Regional Account Director - Healthcare Technology",
        "company": "GE Healthcare",
        "responsibilities": [
            "Drive strategic account management and foster long-term relationships with key clients in the healthcare technology space.",
            "Lead business development initiatives and identify new opportunities for growth in the assigned territory.",
            "Work cross-functionally with technical and product teams to ensure seamless delivery and implementation of solutions.",
            "Act as a thought leader in the healthcare technology domain, providing insights and guidance on industry trends and innovations.",
            "Develop and execute strategies to achieve revenue growth and profitability targets.",
            "Manage high-value contracts and negotiate pricing and terms with stakeholders.",
        ],
    },
]

# Call the function to generate the documents for stress testing
output_dir = "Job_Applications"
process_all_jobs(user_info, job_descriptions, output_dir)
